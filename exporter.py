import re
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Pt, RGBColor, Inches


def parse_markdown_sections(text: str) -> list:
    """Parse markdown into a list of (type, content) tuples."""
    lines = text.split("\n")
    sections = []
    for line in lines:
        line = line.encode('ascii', 'replace').decode('ascii')
        line = line.replace('?', ' ')

        if line.startswith("## "):
            sections.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            sections.append(("h3", line[4:].strip()))
        elif line.startswith("# "):
            sections.append(("h1", line[2:].strip()))
        elif line.strip().startswith("- **") or line.strip().startswith("* **"):
            clean = re.sub(r'^[-*]\s+', '', line.strip())
            sections.append(("bullet_bold", clean))
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            clean = re.sub(r'^[-*]\s+', '', line.strip())
            sections.append(("bullet", clean))
        elif line.strip().startswith("**") and line.strip().endswith("**") and len(line.strip()) > 4:
            sections.append(("bold", line.strip()[2:-2].strip()))
        elif line.strip() == "---":
            sections.append(("hr", ""))
        elif line.strip() == "":
            sections.append(("space", ""))
        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            sections.append(("text", clean.strip()))
    return sections


def export_pdf(report_text: str, topic: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    gold = colors.HexColor("#C8A96E")
    dark = colors.HexColor("#1a1a1a")

    title_style = ParagraphStyle("Title", fontSize=20, textColor=dark,
                              spaceAfter=4, fontName="Times-Roman", 
                              alignment=TA_LEFT, wordWrap='CJK')
    topic_style = ParagraphStyle("Topic", fontSize=16, textColor=gold,
                                  spaceAfter=6, fontName="Times-BoldItalic", alignment=TA_LEFT)
    h2_style = ParagraphStyle("H2", fontSize=14, textColor=dark, spaceBefore=16,
                               spaceAfter=4, fontName="Times-Bold")
    h3_style = ParagraphStyle("H3", fontSize=12, textColor=colors.HexColor("#444444"),
                               spaceBefore=10, spaceAfter=4, fontName="Times-BoldItalic")
    body_style = ParagraphStyle("Body", fontSize=10, textColor=dark, spaceAfter=6,
                                 fontName="Times-Roman", leading=16)
    bold_style = ParagraphStyle("Bold", fontSize=10, textColor=dark, spaceAfter=4,
                                 fontName="Times-Bold")
    bullet_style = ParagraphStyle("Bullet", fontSize=10, textColor=dark, spaceAfter=4,
                                   fontName="Times-Roman", leading=16,
                                   leftIndent=20, firstLineIndent=-12)
    bullet_bold_style = ParagraphStyle("BulletBold", fontSize=10, textColor=dark, spaceAfter=4,
                                        fontName="Times-Roman", leading=16,
                                        leftIndent=20, firstLineIndent=-12)

    story = []

    long_topic = len(topic) > 40
    adjusted_topic_style = ParagraphStyle("TopicAdjusted", fontSize=12 if long_topic else 16,
                                           textColor=gold, spaceAfter=6,
                                           fontName="Times-BoldItalic", alignment=TA_LEFT,
                                           leading=18)
    story.append(Paragraph("Research Report", title_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph(topic.title(), adjusted_topic_style))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=1, color=gold, spaceAfter=20))

    sections = parse_markdown_sections(report_text)
    for kind, content in sections:
        if kind == "h1":
            story.append(Paragraph(content, title_style))
        elif kind == "h2":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                     color=colors.HexColor("#dddddd"), spaceBefore=8))
            story.append(Paragraph(content, h2_style))
        elif kind == "h3":
            story.append(Paragraph(content, h3_style))
        elif kind == "bold":
            story.append(Paragraph(content, bold_style))
        elif kind == "bullet":
            story.append(Paragraph(f"• {content}", bullet_style))
        elif kind == "bullet_bold":
            clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
            story.append(Paragraph(f"• {clean}", bullet_bold_style))
        elif kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                     color=colors.HexColor("#dddddd")))
        elif kind == "space":
            story.append(Spacer(1, 6))
        elif kind == "text" and content.strip():
            story.append(Paragraph(content, body_style))

    doc.build(story)
    return buffer.getvalue()


def export_docx(report_text: str, topic: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading(level=0)
    title.clear()
    run = title.add_run("Research Report")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    topic_para = doc.add_paragraph()
    topic_run = topic_para.add_run(topic.title())
    topic_run.font.size = Pt(16)
    topic_run.font.color.rgb = RGBColor(0xC8, 0xA9, 0x6E)
    topic_run.font.italic = True
    topic_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("─" * 60)

    sections = parse_markdown_sections(report_text)
    for kind, content in sections:
        if kind in ("h1", "h2"):
            h = doc.add_heading(content, level=1 if kind == "h1" else 2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        elif kind == "h3":
            doc.add_heading(content, level=3)
        elif kind == "bold":
            p = doc.add_paragraph()
            r = p.add_run(content)
            r.bold = True
            r.font.size = Pt(11)
        elif kind == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(content)
        elif kind == "bullet_bold":
            p = doc.add_paragraph(style='List Bullet')
            match = re.match(r'\*\*(.*?)\*\*:?\s*(.*)', content)
            if match:
                r = p.add_run(match.group(1) + ": ")
                r.bold = True
                p.add_run(match.group(2))
            else:
                p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', content))
        elif kind == "space":
            doc.add_paragraph()
        elif kind == "text" and content.strip():
            p = doc.add_paragraph(content)
            p.style.font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()